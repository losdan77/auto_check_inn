import glob
import re
import os
import random
import requests
import time
import datetime
import pdfplumber
import shutil
import asyncio
import logging
import pandas as pd
from fastapi import FastAPI
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from bs4 import BeautifulSoup
from docx import Document

from utils import USER_AGENTS, clear_result_folder, SELENIUM_HOST, SELENIUM_PORT


logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


app = FastAPI()


@app.post('/main')
async def main(inn_list: str):
    inn_list = inn_list.split(',')

    number_answer = 0

    for inn in inn_list:
        number_answer += 1
        logger.info(f'Начал {number_answer} итерацию инн:{inn}')
        inn_firm = inn
        
        # 1 -------- файл качается в докер, надо шарить папку и раскоментить парс пдфа с сайта
        try:
            logger.info(f'Начал проход по сайтам')
            start_time = time.time()
            
            first_site_info = first_site(inn = inn, number_answer=number_answer)

            logger.info(f'Конец первого сайта, результаты: {first_site_info}')
            
            array_inn = [inn_firm]
            array_fio = []
            for info in first_site_info[1]:
                array_inn.append(info['inn'])
                array_fio.append(f'{info['surname']} {info['name']} {info['patronymic']}')

            
            second_site_tasks = [asyncio.to_thread(second_site, inn, number_answer) for inn in array_inn]
            fourth_site_tasks = [asyncio.to_thread(fourth_site, inn, number_answer) for inn in array_inn]
            seven_site_tasks = [asyncio.to_thread(seven_site, inn, number_answer) for inn in array_inn]
            # nine_site_tasks = [asyncio.to_thread(nine_site, inn) for inn in array_inn]
            # ten_site_tasks = [asyncio.to_thread(ten_site, fio) for fio in array_fio]

            logger.info(f'Начал асинхронный проход по сайтам')
            (
                second_site_info,
                fourth_site_info,
                seven_site_info,
                # nine_site_info,
                # ten_site_info,
            ) = await asyncio.gather(
                asyncio.gather(*second_site_tasks),
                asyncio.gather(*fourth_site_tasks),
                asyncio.gather(*seven_site_tasks),
                # asyncio.gather(*nine_site_tasks),
                # asyncio.gather(*ten_site_tasks),
            )
            logger.info(f'Закончил асинхронный проход по сайтам')
            # 2 -------- данные есть и скриншот (скриншот как на 7 сайте, надо спросить норм или нет)
            
            # second_site_info = []
            # for inn in array_inn:
            #     info_dict = second_site(inn)
            #     second_site_info.append(info_dict)


            # 3 -------- скриншот есть (единственное спросить хватит только инн фирмы и нужны ли данные в отчет)
            logger.info(f'Начал синхронный проход по сайтам')
            third_site_info = third_site(inn_firm, number_answer) 

            
            # 4 -------- данные и скриншоты есть (по скриншотам надо вопрос задать, что там еще должно быть, может перейти во внутрь карточки)
            # fourth_site_info = []
            # for inn in array_inn:
            #     info_dict = fourth_site(inn)
            #     fourth_site_info.append(info_dict)


            # 5 --------
            # five_site_info = []
            # for inn in array_inn:
            #     info_dict = {f'{inn}': five_site(inn)}
            #     five_site_info.append(info_dict)
            five_site_info = five_site(inn)


            # 7 -------- данные и скриншот есть (скриншот надо спросить пойдет или нет)
            # seven_site_info = []
            # for inn in array_inn:
            #     info_dict = seven_site(inn)
            #     seven_site_info.append(info_dict)


            # 9 -------- файл и данные есть (файл в докере)
            nine_site_info = []
            for inn in array_inn:
                info_dict = nine_site(inn)
                nine_site_info.append(info_dict)
            
            
            # 10 -------- скрины и данные есть
            ten_site_info = []
            for fio in array_fio:
                info_dict = ten_site(fio=fio, number_answer=number_answer)
                ten_site_info.append(info_dict)

            logger.info(f'Закончил синхронный проход по сайтам')
            logger.info(f'Закончил проход по сайтам')

        except Exception as e:
            logger.error(f'Ошибка при прохождении сайтов: {e}')
            time.sleep(5)
            clear_result_folder()
            return 'Запустите еще раз', e

        logger.info(f'Начал формировать ворд')
        document = Document()
        document.add_heading(f'ИНН проверяемой фирмы: {inn_firm} - {first_site_info[0]}')
        
        document.add_paragraph('1) Обращение в ЕГРЮЛ:')
        for i in range(len(array_fio)):
            if i == 0:
                document.add_paragraph(f'{array_fio[i]} - {array_inn[i+1]} - директор')
            else:
                document.add_paragraph(f'{array_fio[i]} - {array_inn[i+1]} - учредители')
        
        document.add_paragraph('2) Обращение в РНП:')
        for info in second_site_info:
            document.add_paragraph(f'{info}')

        document.add_paragraph('4) Обращение в Федресурс:')
        for info in fourth_site_info:
            document.add_paragraph(f'{info}')

        document.add_paragraph('7) Обращение по КОАП:')
        for info in seven_site_info:
            document.add_paragraph(f'{info}')

        document.add_paragraph('9) Обращение в РИА:')
        for info in nine_site_info:
            document.add_paragraph(f'{info}')

        document.add_paragraph('10) Обращение в РЭТ:')
        for info in ten_site_info:
            document.add_paragraph(f'{info}')

        document.save(f'./result/{number_answer}_{inn_firm}.docx')
        logger.info(f'Закончил формировать ворд')
        logger.info(f'Закончил {number_answer} итерацию')
    
    logger.info(f'Закончил прохождение цикла')

    try:
        now_time = str(datetime.datetime.now().strftime("%Y%m%d%H%M"))
        shutil.make_archive(f'{now_time}', 'zip', './result')

        end_time = time.time()
        process_time = end_time - start_time
        logger.info(f'Закончил выполнение работы, времени затрачено: {process_time}')

        return FileResponse(path=f'{now_time}.zip', filename=f'{now_time}.zip') # подумать как удалять его
    
    except Exception as e:
        logger.error(f'Ошибка при отдаче зипника: {e}')
        return 'Запустите еще раз'
    
    finally:
        # os.remove(f'{now_time}.zip')
        clear_result_folder()
        logger.info(f'Очистил папку с результатми')
        


# random_user_agent = random.choice(USER_AGENTS)

headers = {
    'User-Agent': random.choice(USER_AGENTS)
}

download_dir = "./result" 

options = webdriver.ChromeOptions()
# options.add_argument(f"user-agent={random_user_agent}")
options.add_argument("--disable-notifications")  # Отключить уведомления
options.add_argument("--disable-popup-blocking")  # Отключить блокировку всплывающих окон

options.add_experimental_option("prefs", {
    "download.default_directory": download_dir,  # Куда сохранять
    "download.prompt_for_download": False,  # Не показывать окно подтверждения
    "download.directory_upgrade": True,
    "safebrowsing.enabled": True
})


def first_site(inn: str, number_answer: int):
    try:
        logger.info('Начало первого сайта')
        random_user_agent = random.choice(USER_AGENTS)
        options.add_argument(f"user-agent={random_user_agent}")
        driver = webdriver.Remote(
            command_executor=f"http://{SELENIUM_HOST}:{SELENIUM_PORT}/wd/hub",
            options=options  
        )
        driver.get("https://egrul.nalog.ru/index.html")
        time.sleep(5)
        input_inn = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//input[@id='query']"))
        )
        print(input_inn.get_attribute('outerHTML'))
        
        input_inn.click()
        time.sleep(1)

        input_inn.send_keys(inn)
        input_inn.send_keys(Keys.RETURN)
        time.sleep(15)
        # result_title = WebDriverWait(driver, 20).until(
        #     EC.presence_of_element_located((By.XPATH, "//div[contains(@class, 'res-caption')]"))
        # ) 
        # result_info = WebDriverWait(driver, 20).until(
        #     EC.presence_of_element_located((By.XPATH, "//div[contains(@class, 'res-text')]"))
        # )
        button_for_download_pdf = WebDriverWait(driver, 20).until(
            EC.presence_of_element_located((By.XPATH, "//button[contains(@class, 'btn-with-icon btn-excerpt op-excerpt')]"))
        )
        
        button_for_download_pdf.click()
        time.sleep(25)
        pdf_file = glob.glob("./result/ul-*.pdf")
        
        with pdfplumber.open(f'{pdf_file[0]}') as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        
        firm_name_pattern = r"1 Полное наименование на русском языке(.*?)2"
        firm_name_match = re.search(firm_name_pattern, text, re.DOTALL)

        firm_name = firm_name_match.group(1).strip().replace("\n", " ").replace("\"", "'") if firm_name_match else "Не найдено"
        
    
        all_partition_pattern = r"\d+\s+Фамилия\s+([^\n]+)\s+Имя\s+([^\n]+)\s+Отчество\s+([^\n]+)\s+\d+\s+ИНН\s+(\d+)"
        all_partition_matches = re.findall(all_partition_pattern, text)
        all_partition_result = []
        
        for match in all_partition_matches:
            surname, name, patronymic, inn_lica = match
            result_dict = {
                'surname': surname,
                'name': name,
                'patronymic': patronymic,
                'inn': inn_lica,
            }
            all_partition_result.append(result_dict)
        
        logger.info(f'Конец первого сайта, результаты: {firm_name, all_partition_result}')
        return firm_name, all_partition_result
    
    except Exception as e:
        logger.error(f'Ошибка в первом сайте: {e}')
    finally:
        # переменовать пдф
        if os.path.isfile(pdf_file[0]) or os.path.islink(pdf_file[0]):
            new_pdf_file = pdf_file[0].replace(f'{pdf_file[0]}', f'./result/{number_answer}.pdf')
            os.rename(pdf_file[0], new_pdf_file)

        driver.quit()
        


def second_site(inn: str, number_answer: int):
    try:
        logger.info('Начало второго сайта')
    
        url = f"https://zakupki.gov.ru/epz/dishonestsupplier/search/results.html?searchString={inn}&morphology=on&sortBy=UPDATE_DATE&pageNumber=1&sortDirection=false&recordsPerPage=_10&showLotsInfoHidden=false&fz94=on&fz223=on&ppRf615=on"
        
        response = requests.get(url, headers=headers)

        soup = BeautifulSoup(response.text, "lxml")
        result_info = soup.find(class_='search-registry-entrys-block')

        ##############################
        random_user_agent = random.choice(USER_AGENTS)
        options.add_argument(f"user-agent={random_user_agent}")
        driver = webdriver.Remote(
            command_executor=f"http://{SELENIUM_HOST}:{SELENIUM_PORT}/wd/hub",
            options=options
        )
        driver.get(f"https://zakupki.gov.ru/epz/dishonestsupplier/search/results.html?searchString={inn}&morphology=on&sortBy=UPDATE_DATE&pageNumber=1&sortDirection=false&recordsPerPage=_10&showLotsInfoHidden=false&fz94=on&fz223=on&ppRf615=one")
        
        time.sleep(5)
        driver.save_screenshot(f'./result/screenshot/second_site/{number_answer}_{inn}.png')
        #############################
        
        logger.info('Конец второго сайта')
        if len(result_info.text) > 10:
            return {f'{inn}': 'да'}
        return {f'{inn}': 'нет'}
    
    except Exception as e:
        logger.error(f'Ошибка во втором сайте: {e}')
    finally:
        driver.quit()


def third_site(inn: str, number_answer: int):
    try:
        logger.info('Начало третьего сайта')

        random_user_agent = random.choice(USER_AGENTS)
        options.add_argument(f"user-agent={random_user_agent}")
        driver = webdriver.Remote(
            command_executor=f"http://{SELENIUM_HOST}:{SELENIUM_PORT}/wd/hub",
            options=options  
        )
        driver.get(f"https://pb.nalog.ru/index.html")

        input_inn = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//input[contains(@class, 'u3-editor')]"))
        )
       
        input_inn.send_keys(inn)
        input_inn.send_keys(Keys.RETURN)
        
        time.sleep(5)
        driver.save_screenshot(f'./result/screenshot/third_site/{number_answer}_{inn}(1).png')

        name_organization = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//div[contains(@class, 'pb-card__title')]"))
        )
        name_organization.click()

        time.sleep(5)
        result_info = driver.find_element(By.XPATH, "//span[contains(text(), 'Сведения о лице, имеющем право без доверенности действовать от имени юридического лица')]")
        ActionChains(driver).scroll_to_element(result_info).perform()
        driver.save_screenshot(f'./result/screenshot/third_site/{number_answer}_{inn}(2).png')
        
        time.sleep(2)
        result_info = driver.find_element(By.XPATH, "//span[contains(text(), 'Сведения о непредставлении налоговой отчетности более года')]")
        ActionChains(driver).scroll_to_element(result_info).perform()
        driver.save_screenshot(f'./result/screenshot/third_site/{number_answer}_{inn}(3).png')

        logger.info('Конец третьего сайта')

    except Exception as e:
        logger.error(f'Ошибка в третьем сайте: {e}')
    finally:
        driver.quit()
        


def fourth_site(inn: str, number_answer: int):
    try:
        logger.info('Начало четвертого сайта')

        random_user_agent = random.choice(USER_AGENTS)
        options.add_argument(f"user-agent={random_user_agent}")
        driver = webdriver.Remote(
            command_executor=f"http://{SELENIUM_HOST}:{SELENIUM_PORT}/wd/hub",
            options=options  
        )
        driver.get(f"https://fedresurs.ru/entities?searchString={inn}")
        
        time.sleep(5)
        
        result = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//div[contains(@class, 'container full-height')]"))
        )

        driver.save_screenshot(f'./result/screenshot/fourth_site/{number_answer}_{inn}.png')
        
        logger.info('Конец четвертого сайта')
        if len(result.get_attribute('outerHTML')) > 7000:
            return {f'{inn}': 'да'}
        return {f'{inn}': 'нет'}
    
    except Exception as e:
        logger.error(f'Ошибка в четвертом сайте: {e}')
    finally:
        driver.quit()


def five_site(inn: str):
    try:
        logger.info('Начало пятого сайта')
        # inn = "5054004240"
        # proxy = "http://V84kEe:XhAdiJu5Ej@45.15.72.224:5500"
        # options.add_argument(f"--proxy-server={proxy}")
        random_user_agent = random.choice(USER_AGENTS)
        options.add_argument(f"user-agent={random_user_agent}")
        driver = webdriver.Remote(
            command_executor=f"http://{SELENIUM_HOST}:{SELENIUM_PORT}/wd/hub",
            options=options  
        )
        
        driver.get(f"https://kad.arbitr.ru/")

        time.sleep(5)
        driver.execute_script("""
            var modal = document.querySelector('.b-browsers.js-browsers');
            if (modal) {
                modal.style.display = 'none';
            }
        """)
        
        # Попробовать закрыть всплывающее окно
        try:
            # close_button = driver.find_element(By.XPATH, "//div[contains(@class='b-promo_notification-popup_wrapper')]//a[contains(@class='b-promo_notification-popup-close js-promo_notification-popup-close')]")
            close_button = driver.find_element(By.XPATH, "//a[contains(@class,'b-promo_notification-popup-close js-promo_notification-popup-close')]")
            close_button.click()
        except Exception as e:
            print("Не удалось закрыть всплывающее окно:", e)
        driver.save_screenshot(f'./result/screenshot/five_site/01_{inn}.png')
        time.sleep(2)

        input_inn = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//textarea[contains(@class, 'g-ph')]"))
        )
        print(input_inn.get_attribute('outerHTML'))

        # input_inn.click()
        time.sleep(1)

        input_inn.send_keys(inn)
        driver.save_screenshot(f'./result/screenshot/five_site/02_{inn}.png')
        input_inn.send_keys(Keys.RETURN)
        driver.save_screenshot(f'./result/screenshot/five_site/1_{inn}.png')
        
        
        button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//button[contains(text(), 'Найти')]"))
        )
        print(button.get_attribute('outerHTML'))
        button.click()
        driver.save_screenshot(f'./result/screenshot/five_site/03_{inn}.png')
        time.sleep(10)
        # result_info = WebDriverWait(driver, 10).until( # не получается увидеть таблицу с результатами очень сложный сайт (нужно подумать нужна ли эта переменная)
        #     # EC.element_to_be_clickable((By.XPATH, "//div[@id='table']"))
        #     EC.element_to_be_clickable((By.XPATH, "//div[contains(@class, 'b-results')]"))
        # )
        driver.save_screenshot(f'./result/screenshot/five_site/2_{inn}.png')
        # banckrot_button = WebDriverWait(driver, 10).until(
        #     EC.element_to_be_clickable((By.XPATH, "//li[contains(@class, 'bankruptcy')]"))
        # )
        # banckrot_button.click() # !почему то говори, что не кликабле, хотя по факту кликабле
        time.sleep(2)
        # driver.save_screenshot(f'./result/screenshot/five_site/{inn}.png')
        # print(result_info.get_attribute('outerHTML'))
        

        logger.info('Конец пятого сайта')

    except Exception as e:
        logger.error(f'Ошибка в пятом сайте: {e}')
    finally:
        driver.quit()


def seven_site(inn: str, number_answer: int):
    try:
        logger.info('Начало седьмого сайта')

        url = f'https://zakupki.gov.ru/epz/main/public/document/search.html?searchString={inn}&sectionId=2369&strictEqual=false'

        response = requests.get(url, headers=headers)

        soup = BeautifulSoup(response.text, "lxml")
        result_info = soup.find_all('section')[0].find(class_='docs-list')
        
        print(len(result_info.text))
        ##############################
        random_user_agent = random.choice(USER_AGENTS)
        options.add_argument(f"user-agent={random_user_agent}")
        driver = webdriver.Remote(
            command_executor=f"http://{SELENIUM_HOST}:{SELENIUM_PORT}/wd/hub",
            options=options  
        )
        driver.get(f"https://zakupki.gov.ru/epz/main/public/document/search.html?searchString={inn}&sectionId=2369&strictEqual=false")
        time.sleep(5)
        driver.save_screenshot(f'./result/screenshot/seven_site/{number_answer}_{inn}.png')
        #############################
        
        logger.info('Конец седьмого сайта')
        if len(result_info.text) > 100:
            return {f'{inn}': 'да'}
        return {f'{inn}': 'нет'}
    
    except Exception as e:
        logger.error(f'Ошибка в седьмом сайте: {e}')
    finally:
        driver.quit()


def nine_site(inn: str): # Впринципе готово (проблема по MIMA)
 
    logger.info('Начало девятого сайта')

    excel_file = glob.glob("./result/export.xlsx")
    if not excel_file:
        try:
            logger.info('Качаем эксель девятого сайта')
            url = 'https://minjust.gov.ru/ru/pages/reestr-inostryannykh-agentov/'
            
            random_user_agent = random.choice(USER_AGENTS)
            options.add_argument(f"user-agent={random_user_agent}")
            driver = webdriver.Remote(
                command_executor=f"http://{SELENIUM_HOST}:{SELENIUM_PORT}/wd/hub",
                options=options  
            )
            driver.get(url)

            time.sleep(10)

            button_for_download_excel = WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.XPATH, "//span[contains(text(), 'Загрузить реестр')]"))
            )
            
            button_for_download_excel.click()
            time.sleep(20)
            logger.info('Закончили качать эксель девятого сайта')

        except Exception as e:
            logger.error(f'Ошибка в скачивании экселя с девятого сайта: {e}')
        finally:
            driver.quit()

    excel_path = "./result/export.xlsx"
    target_inn = inn

    try:
        # Читаем Excel-файл
        df = pd.read_excel(excel_path, engine='openpyxl')
        
        # Проверяем, достаточно ли столбцов в файле
        if df.shape[1] < 9:
            print("Файл содержит меньше 9 столбцов, столбец 'I' отсутствует.")
            exit()
        
        # Получаем данные из столбца 'I' (девятый столбец, индекс 8)
        column_i = df.iloc[:, 8].astype(str)
        
        # Проверяем наличие ИНН в столбце 'I'
        if target_inn in column_i.values:
            return {f'{inn}': 'присутствует в реестре'}
        else:
            return {f'{inn}': 'нет'}

    except FileNotFoundError:
        logger.error('Файл девятого сайта не найден')
    except Exception as e:
        logger.error(f'Ошибка при парсинге экселя девятого сайта: {e}')



def ten_site(fio: str, number_answer: int):
    try:
        logger.info('Начало десятого сайта')

        random_user_agent = random.choice(USER_AGENTS)
        options.add_argument(f"user-agent={random_user_agent}")
        driver = webdriver.Remote(
            command_executor=f"http://{SELENIUM_HOST}:{SELENIUM_PORT}/wd/hub",
            options=options  
        )
        driver.get(f"https://www.fedsfm.ru/documents/terr-list")

        time.sleep(5)

        form_for_fio = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//input[contains(@class, 'form-control')]"))
        )
        
        form_for_fio.send_keys(fio)
        form_for_fio.send_keys(Keys.RETURN)

        time.sleep(5)
        try:
            WebDriverWait(driver, 5).until(EC.alert_is_present())
            alert = driver.switch_to.alert
            print(f"Alert text: {alert.text}")
            alert.accept()  # Закрыть alert
        except:
            pass
        time.sleep(5)

        result_info = WebDriverWait(driver, 15).until(
            EC.element_to_be_clickable((By.XPATH, "//tbody"))
        )
        driver.save_screenshot(f'./result/screenshot/ten_site/{number_answer}_{fio}.png')
    
        logger.info('Конец десятого сайта')
        if len(result_info.get_attribute('outerHTML')) > 120:
            return {f'{fio}': 'да'}
        return {f'{fio}': 'нет'}
    
    except Exception as e:
        logger.error(f'Ошибка в десятом сайте: {e}')
        return {f'{fio}': 'НЕ ИЗВЕСТНО (Необходимо проверить вручную)'}

    finally:
        driver.quit()
